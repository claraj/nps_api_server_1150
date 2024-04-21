import docx
import requests

base_server = 'http://127.0.0.1:5000'
base_server = 'http://national-parks-1150.azurewebsites.net'

url = base_server + '/api/'

parks = requests.get(url + 'list').json()

doc = docx.Document()

# problems = ['GLAC', 'ROMO', 'ZION']
# parks = [ { 'park_code': problem} for problem in problems ]
# print("HELLO??", parks)
for park in parks:
    print(park)

    park_resp = requests.get(url + park['park_code']).json()
    imgs = park_resp['images']
    for img in imgs:
        resp = requests.get(img['url'])
        if resp.status_code != 200:
            print('Image DOWNLOAD error,', resp.status_code, park['park_code'], img)

        with open('tmp_img/image.jpg', 'wb') as f:
            for chunk in resp:
                f.write(chunk)

        # try:
        doc.add_paragraph(park['park_code'])
        doc.add_paragraph(img['url'])
        doc.add_picture('tmp_img/image.jpg')
        # except Exception as e:
        #     print(e)
        #     print('Error WRITING image', park['park_code'], img, e)

doc.save('tmp_img/worddoc.docx')
       